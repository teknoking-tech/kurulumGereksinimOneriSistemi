"""
CBOT Kurulum Asistanı - Doküman Oluşturma İşlemleri
"""

import os
import tempfile
from datetime import datetime

# PDF oluşturma kütüphanesi
from fpdf import FPDF
from fpdf.enums import Align  # Align enum'ını import ediyoruz

# Word belgesi oluşturma kütüphanesi
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Yardımcı fonksiyonlar
from utils import get_environment_label, get_deployment_label, get_service_label


def generate_markdown_content(form_data, requirements):
    """Markdown içeriği oluşturan yardımcı fonksiyon"""
    
    # Kurum/Şirket bilgisini alma
    org_name = form_data["organization_name"] if form_data["organization_name"] else "Belirtilmemiş"
    
    content = f"""# CBOT Kurulum Gereksinimleri

## 1. Proje Bilgileri
- **Kurum/Şirket:** {org_name}
- **Proje Kodu:** {form_data["project_code"] if form_data["project_code"] else "Belirtilmemiş"}
- **İletişim E-posta:** {form_data["contact_email"] if form_data["contact_email"] else "Belirtilmemiş"}
- **Ortam Tipi:** {', '.join([get_environment_label(e) for e in form_data["environment_type"]])}
- **Deployment Tipi:** {get_deployment_label(form_data["deployment_type"])}

## 2. Seçilen Modüller ve Servisler

### Ana Modüller:
"""

    # Seçilen modülleri ekleyin
    for module_name, is_selected in form_data["modules"].items():
        if is_selected:
            content += f"- {module_name.capitalize()}\n"

    content += """
### Yardımcı Servisler:
"""
    
    # Seçilen servisleri ekleyin
    services_selected = False
    for service_name, is_selected in form_data["auxiliary_services"].items():
        if is_selected:
            services_selected = True
            content += f"- {get_service_label(service_name)}\n"
    
    if not services_selected:
        content += "- Seçilen yardımcı servis bulunmamaktadır.\n"

    content += """
## 3. Donanım Gereksinimleri

"""
    
    # Donanım gereksinimlerini ekleyin
    if requirements["hardware"] and len(requirements["hardware"]) > 0:
        for module, env_reqs in requirements["hardware"].items():
            content += f"### {module.capitalize()} Modülü\n"
            
            if "test" in form_data["environment_type"] and "test" in env_reqs:
                content += f"- **Test:** {env_reqs['test']['cpu']}, {env_reqs['test']['ram']}, {env_reqs['test']['disk']} disk\n"
            
            if "live" in form_data["environment_type"] and "live" in env_reqs:
                content += f"- **Canlı:** {env_reqs['live']['cpu']}, {env_reqs['live']['ram']}, {env_reqs['live']['disk']} disk\n"
    else:
        content += "**Seçilen modül bulunamadı**\n"

    content += """
### İşletim Sistemi
- CentOS 7/8 veya RHEL

## 4. Veritabanı Gereksinimleri
"""
    
    # Veritabanı gereksinimlerini ekleyin
    if "database" in requirements and requirements["database"]:
        content += f"- **Tip:** {requirements['database'].get('type', 'Belirtilmemiş')}\n"
        content += f"- **Port:** {requirements['database'].get('port', 'Belirtilmemiş')}\n"
        
        if requirements['database'].get('collation') and requirements['database']['collation'] != 'N/A':
            content += f"- **Collation:** {requirements['database']['collation']}\n"
        
        content += f"- **Kullanıcı Yetkisi:** {requirements['database'].get('permissions', 'Belirtilmemiş')}\n"
    else:
        content += "- Veritabanı gereksinimleri belirtilmemiş\n"

    content += """
## 5. Network / Firewall Gereksinimleri

### DNS Kayıtları
"""
    
    # DNS kayıtlarını ekleyin
    if "network" in requirements and "dns_records" in requirements["network"] and requirements["network"]["dns_records"]:
        for record in requirements["network"]["dns_records"]:
            note = f" ({record['note']})" if "note" in record else ""
            content += f"- **{record['name']}** -> Port: {record['port']}{note}\n"
    else:
        content += "- Kayıt bulunamadı\n"

    content += """
### Genel Ağ Gereksinimleri
- Sunucular internet erişimli olmalıdır
- Docker imajlarının çekilebilmesi için registry.cbot.ai adresine 443 portundan erişim sağlanmalıdır
"""
    
    # Veritabanı portu varsa ekleyin
    db_port = requirements["database"].get("port", "?") if "database" in requirements else "?"
    content += f"- Veritabanı portu ({db_port}) tüm modüllerden erişilebilir olmalıdır\n"
    content += "- Load Balancer kullanılıyorsa, a5000 portunda WebSocket desteği sağlanmalıdır\n"

    # LDAP gereksinimleri bölümü
    section_num = 6
    if form_data["use_ldap"] and "ldap" in requirements and requirements["ldap"]:
        content += f"""
## {section_num}. LDAP Gereksinimleri

### Gerekli Alanlar:
"""
        for field in requirements["ldap"]["fields"]:
            content += f"- {field}\n"
        
        content += """
### Alan Eşleştirmeleri:
"""
        for field in requirements["ldap"]["mapping"]:
            content += f"- {field}\n"
        
        content += f"\n**Not:** {requirements['ldap']['note']}\n"
        section_num += 1

    # Ek gereksinimler ve uyarılar
    content += f"""
## {section_num}. Ek Gereksinimler ve Uyarılar
"""
    
    if "additional_requirements" in requirements and requirements["additional_requirements"]:
        for req in requirements["additional_requirements"]:
            content += f"- {req}\n"
    else:
        content += "- Ek gereksinim bulunmamaktadır.\n"

    # Son bölüm
    current_date = datetime.now().strftime("%d.%m.%Y")
    content += f"""
---

*Bu doküman CBOT Kurulum Asistanı v2.0 tarafından {current_date} tarihinde otomatik olarak oluşturulmuştur.*
"""
    
    return content


def create_markdown_file(form_data, requirements):
    """Markdown dosyası oluşturun"""
    content = generate_markdown_content(form_data, requirements)
    
    # Geçici bir dosya oluşturun
    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, "cbot_requirements.md")
    
    with open(file_path, "w", encoding="utf-8") as file:
        file.write(content)
    
    return file_path


def create_pdf_file(form_data, requirements):
    """PDF dosyası oluşturun"""
    # Markdown içeriğini alın
    markdown_content = generate_markdown_content(form_data, requirements)
    
    # Basit bir FPDF belgesi oluşturma
    pdf = FPDF()
    pdf.add_page()
    
    # Başlık formatı
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, "CBOT Kurulum Gereksinimleri", 0, 1, "C")
    pdf.ln(5)
    
    # İçerik formatı - Basitleştirilmiş sürüm
    # Not: Gerçek bir uygulamada Markdown'dan PDF'e daha gelişmiş bir dönüşüm gerekebilir
    pdf.set_font("Arial", "", 11)
    
    # Satır satır içeriği işleyin (basit yaklaşım)
    for line in markdown_content.split("\n"):
        if line.startswith("# "):
            pdf.set_font("Arial", "B", 14)
            pdf.ln(5)
            pdf.cell(0, 10, line[2:], 0, 1)
            pdf.ln(2)
            pdf.set_font("Arial", "", 11)
        elif line.startswith("## "):
            pdf.set_font("Arial", "B", 12)
            pdf.ln(3)
            pdf.cell(0, 8, line[3:], 0, 1)
            pdf.ln(2)
            pdf.set_font("Arial", "", 11)
        elif line.startswith("### "):
            pdf.set_font("Arial", "B", 11)
            pdf.ln(2)
            pdf.cell(0, 6, line[4:], 0, 1)
            pdf.ln(1)
            pdf.set_font("Arial", "", 11)
        elif line.startswith("- "):
            pdf.cell(5, 5, "*", 0, 0)
            pdf.multi_cell(0, 5, line[2:])
        elif line.startswith("**Not:**"):
            pdf.set_font("Arial", "I", 10)
            pdf.multi_cell(0, 5, line)
            pdf.set_font("Arial", "", 11)
        elif line == "---":
            pdf.ln(5)
            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(5)
        elif line.startswith("*Bu doküman"):
            pdf.set_font("Arial", "I", 9)
            pdf.multi_cell(0, 5, line)
        elif line:
            pdf.multi_cell(0, 5, line)
    
    # Geçici bir dosya oluşturun
    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, "cbot_requirements.pdf")
    
    pdf.output(file_path)
    return file_path


def create_word_file(form_data, requirements):
    """Word belgesi oluşturun"""
    # Markdown içeriğini alın
    markdown_content = generate_markdown_content(form_data, requirements)
    
    # Word belgesi oluşturma
    doc = Document()
    
    # Sayfa kenar boşlukları
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    # Başlık formatı
    title = doc.add_heading("CBOT Kurulum Gereksinimleri", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # İçerik formatı - Basitleştirilmiş sürüm
    # Not: Gerçek bir uygulamada Markdown'dan Word'e daha gelişmiş bir dönüşüm gerekebilir
    
    # Satır satır içeriği işleyin (basit yaklaşım)
    for line in markdown_content.split("\n"):
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            para = doc.add_paragraph()
            para.add_run("• ").bold = True
            para.add_run(line[2:])
        elif line.startswith("**Not:**"):
            # Intense Quote stilini kullanmak yerine özel bir paragraf oluşturma
            para = doc.add_paragraph()
            para.add_run(line).italic = True
        elif line == "---":
            doc.add_paragraph("_" * 50)
        elif line.startswith("*Bu doküman"):
            # Subtle Emphasis stilini kullanmak yerine özel bir paragraf oluşturma
            para = doc.add_paragraph()
            para.add_run(line).italic = True
        elif line:
            doc.add_paragraph(line)
    
    # Geçici bir dosya oluşturun
    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, "cbot_requirements.docx")
    
    doc.save(file_path)
    return file_path


def generate_document(form_data, requirements, format_type):
    """Belirtilen formatta belge oluşturun"""
    if format_type == "markdown":
        return create_markdown_file(form_data, requirements)
    elif format_type == "pdf":
        return create_pdf_file(form_data, requirements)
    elif format_type == "word":
        return create_word_file(form_data, requirements)
    else:
        raise ValueError(f"Desteklenmeyen belge formatı: {format_type}")